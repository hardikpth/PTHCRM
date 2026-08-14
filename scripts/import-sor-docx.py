#!/usr/bin/env python3
"""Import the approved PTH SOR Word tables into the CRM dataset."""

import json
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT.parent / "PTH_SOR_2026-27_Final_Complete.docx"
OUTPUT = ROOT / "assets" / "js" / "sor.js"

TEXT_FIXES = {
    "Compressive Strength (Compete 3, 7 and 28 days cycle)": "Compressive Strength (Complete 3, 7 and 28 days cycle)",
    "Combined Ferric Oxide and Alumina (R203)": "Combined Ferric Oxide and Alumina (R2O3)",
    "Chloride (CI) Content": "Chloride (Cl) Content",
    "Water Soluble Sulphate as SO3 (CI.10)": "Water Soluble Sulphate as SO3 (Cl. 10)",
    "Water Soluble Sulphate as SO4 (CI.10)": "Water Soluble Sulphate as SO4 (Cl. 10)",
    "Sulphate (SO) Content": "Sulphate (SO3) Content",
}


def clean(value):
    return " ".join(value.split())


def numeric_rate(value):
    match = re.search(r"(?<!\d)(\d[\d,]*)(?!\d)", value)
    return int(match.group(1).replace(",", "")) if match else None


def category_titles(document):
    titles = {}
    combos = {}
    current = None
    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        match = re.match(r"^(\d+)\.\s+(.+)$", text)
        if match:
            current = int(match.group(1))
            titles[current] = match.group(2)
        elif current and "COMBO PACKAGE OFFER" in text:
            combos.setdefault(current, []).append(text)
    return titles, combos


def import_table(table, category_id):
    rows = table.rows[1:]
    last_cells = [row.cells[-1]._tc for row in rows]
    is_package = len(rows) > 1 and len({id(cell) for cell in last_cells}) == 1
    package_rate = numeric_rate(rows[0].cells[-1].text) if is_package else None
    tests = []

    for row in rows:
        cells = [clean(cell.text) for cell in row.cells]
        if category_id in (18, 22):
            context, name, code = cells[1], cells[2], cells[3]
            qty = "Not specified"
            rate_text = cells[4]
        elif category_id in (33, 34):
            name, code, rate_text = cells[1], cells[2], cells[3]
            context = ""
            qty = "Per point" if re.search(r"/\s*Point\b", rate_text, re.I) else "As per site scope"
        else:
            name, code, qty, rate_text = cells[1], cells[2], cells[3], cells[4]
            context = ""

        name = TEXT_FIXES.get(name, name)
        if context:
            name = f"{name} — {context}"

        if is_package:
            rate = None
            display_rate = f"Included in package ₹{package_rate:,}"
        elif re.search(r"on\s+(request|demand)", rate_text, re.I):
            rate = None
            display_rate = "Rate on Demand" if "demand" in rate_text.lower() else "On request"
        else:
            rate = numeric_rate(rate_text)
            display_rate = rate_text or (str(rate) if rate is not None else "On request")

        tests.append({
            "name": name,
            "code": code,
            "qty": qty,
            "rate": rate,
            "rateText": display_rate,
        })

    return tests, package_rate


def main():
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE
    if not source.is_file():
        raise SystemExit(f"SOR source not found: {source}")

    document = Document(source)
    if len(document.tables) != 34:
        raise SystemExit(f"Expected 34 SOR tables; found {len(document.tables)}")

    titles, combo_map = category_titles(document)
    categories = []
    for index, table in enumerate(document.tables, 1):
        tests, package_rate = import_table(table, index)
        combos = combo_map.get(index, [])
        if package_rate:
            combos.append(f"PACKAGE RATE: All {len(tests)} listed parameters — ₹{package_rate:,} excluding GST")
        category = {
            "id": index,
            "name": titles[index],
            "combos": combos,
            "tests": tests,
        }
        if package_rate:
            category["packageRate"] = package_rate
        categories.append(category)

    count = sum(len(category["tests"]) for category in categories)
    if count != 310:
        raise SystemExit(f"Expected 310 tests; found {count}")

    meta = {
        "financialYear": "2026-27",
        "gstExclusive": True,
        "gstRate": 18,
        "sourceFile": source.name,
    }
    content = (
        "/* PTH CRM — Schedule of Rates FY 2026-27 (imported from the approved Word SOR) */\n"
        f"window.SOR_META = {json.dumps(meta, ensure_ascii=False, separators=(',', ':'))};\n"
        f"window.SOR = {json.dumps(categories, ensure_ascii=False, separators=(',', ':'))};\n"
    )
    OUTPUT.write_text(content, encoding="utf-8")
    print(f"Imported {count} tests across {len(categories)} categories from {source.name}")


if __name__ == "__main__":
    main()
